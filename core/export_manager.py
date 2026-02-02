#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
导出管理器 - 核心业务逻辑
集成原项目的多格式导出功能
"""

import os
import json
import time
from datetime import datetime
import base64
import re
from io import BytesIO
from .enterprise_logger import app_logger
from .common import get_question_field

# 导入各种格式的导出器
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import xlsxwriter
    XLSXWRITER_AVAILABLE = True
except ImportError:
    xlsxwriter = None
    XLSXWRITER_AVAILABLE = False

EXCEL_AVAILABLE = XLSXWRITER_AVAILABLE

class ExportManager:
    """导出管理器类"""
    
    def __init__(self):
        """初始化导出管理器"""
        self.supported_formats = {
            'markdown': self.export_markdown,
            'html': self.export_html,
            'json': self.export_json,
            'docx': self.export_docx if DOCX_AVAILABLE else None,
            'pdf': self.export_pdf if PDF_AVAILABLE else None,
            'excel': self.export_excel if EXCEL_AVAILABLE else None
        }

    def _resolve_img_src(self, img: dict) -> str:
        """尽可能从多种键中解析图片src，兼容不同解析器输出。
        优先 data(dataURL) → base64 → src/url。
        """
        if not img:
            return ''
        src = img.get('data')
        if not src:
            b64 = img.get('base64')
            if b64:
                # 默认按png封装
                src = f"data:image/png;base64,{b64}"
        if not src:
            src = img.get('src') or img.get('url') or ''
        return src
        
    def export_questions(self, questions_data, format_name, config):
        """导出题目数据"""
        try:
            if format_name not in self.supported_formats:
                raise Exception(f"不支持的导出格式: {format_name}")
                
            exporter = self.supported_formats[format_name]
            if not exporter:
                raise Exception(f"导出格式 {format_name} 不可用，缺少相关依赖")
                
            # 生成文件名
            filename = self.generate_filename(format_name, config)
            
            # 确保保存目录存在
            save_path = config.get('save_path', 'exports/')
            os.makedirs(save_path, exist_ok=True)
            
            full_path = os.path.join(save_path, filename)
            
            # 执行导出
            exporter(questions_data, full_path, config)
            
            return full_path
            
        except Exception as e:
            app_logger.error(f"导出失败: {e}")
            raise
            
    def generate_filename(self, format_name, config):
        """生成文件名"""
        pattern = config.get('filename_pattern', 'exported_questions_{timestamp}')
        
        # 替换变量
        add_ts = config.get('add_timestamp', True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S') if add_ts else ''
        filename = pattern.replace('{timestamp}', timestamp)
        filename = filename.replace('{format}', format_name)
        
        # 添加扩展名
        extensions = {
            'markdown': '.md',
            'html': '.html',
            'json': '.json',
            'docx': '.docx',
            'pdf': '.pdf',
            'excel': '.xlsx'
        }
        
        if not filename.endswith(extensions.get(format_name, '')):
            filename += extensions.get(format_name, '')
            
        return filename
        
    def export_markdown(self, questions_data, filepath, config):
        """导出为Markdown格式"""
        try:
            # 流式写入，降低内存占用
            f = open(filepath, 'w', encoding='utf-8')
            def write_line(line: str = ""):
                f.write(line + "\n")
            
            # 标题
            title = config.get('custom_title') or "题目导出"
            write_line(f"# {title}\n")
            
            # 元信息
            write_line(f"**导出时间**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
            write_line(f"**题目总数**: {len(questions_data)}\n")
            
            # 按配置分组
            grouped_questions = self.group_questions(questions_data, config)
            
            for group_name, questions in grouped_questions.items():
                if group_name != "default":
                    write_line(f"## {group_name} ({len(questions)}题)\n")
                    
                for i, question in enumerate(questions, 1):
                    write_line(self.format_question_markdown(question, i, config))
                    write_line("---\n")
                    
            # 统计信息
            if config.get('include_statistics', True):
                write_line(self.generate_statistics_markdown(questions_data))
                
            f.close()
                
        except Exception as e:
            raise Exception(f"Markdown导出失败: {e}")
            
    def format_question_markdown(self, question, index, config):
        """格式化单个题目为Markdown"""
        lines = []
        
        # 题目标题
        question_text = question.get('title', question.get('question_text', ''))
        question_type = question.get('question_type', '')
        if question_type:
            lines.append(f"### {index}. ({question_type}){question_text}")
        else:
            lines.append(f"### {index}. {question_text}")
        
        # 题干图片
        title_images = question.get('title_images') or []
        if config.get('include_images', True) and title_images:
            for img in title_images:
                img_src = self._resolve_img_src(img)
                if img_src:
                    lines.append(f"\n![{img.get('alt','图片')}]({img_src})")
            
        # 选项
        if config.get('include_options', True) and question.get('options'):
            lines.append("\n**选项**:")
            for option in question['options']:
                label = option.get('label', '')
                content = option.get('content', '')
                opt_images = option.get('images') or []
                
                # 如果有图片，清理占位符文本
                if opt_images:
                    content = re.sub(r'\[图片选项[：:]\s*\d+张\]', '', content).strip()
                
                lines.append(f"- {label}. {content}")
                # 选项图片
                if config.get('include_images', True) and opt_images:
                    for img in opt_images:
                        img_src = self._resolve_img_src(img)
                        if img_src:
                            lines.append(f"  \n  ![{img.get('alt','图片')}]({img_src})")
                
        # 答案
        if config.get('include_answers', True):
            my_answer = question.get('my_answer', '')
            correct_answer = question.get('correct_answer', '')
            if my_answer:
                lines.append(f"\n**我的答案**: {my_answer}")
            if correct_answer:
                lines.append(f"**正确答案**: {correct_answer}")
                
        # 得分
        score = question.get('score', '')
        if score:
            lines.append(f"**得分**: {score}")
            
        # 解析
        if config.get('include_explanations', True):
            explanation = question.get('explanation', '')
            if explanation:
                lines.append(f"\n**解析**: {explanation}")
                # 解析图片
                if config.get('include_images', True):
                    for img in question.get('explanation_images') or []:
                        img_src = self._resolve_img_src(img)
                        if img_src:
                            lines.append(f"\n![{img.get('alt','图片')}]({img_src})")
                
        return '\n'.join(lines)
        
    def export_html(self, questions_data, filepath, config):
        """导出为HTML格式"""
        try:
            title = config.get('custom_title') or "题目导出"
            
            html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        {self.get_html_styles()}
    </style>
</head>
<body>
    <div class="container">
        <header class="header">
            <h1>{title}</h1>
            <div class="meta">
                <span>导出时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}</span>
                <span>题目总数: {len(questions_data)}</span>
            </div>
        </header>
        
        <main class="content">
"""
            
            # 按配置分组
            grouped_questions = self.group_questions(questions_data, config)
            
            for group_name, questions in grouped_questions.items():
                if group_name != "default":
                    html_content += f'<section class="group"><h2>{group_name} ({len(questions)}题)</h2>'
                    
                for i, question in enumerate(questions, 1):
                    html_content += self.format_question_html(question, i, config)
                    
                if group_name != "default":
                    html_content += '</section>'
                    
            # 统计信息
            if config.get('include_statistics', True):
                html_content += self.generate_statistics_html(questions_data)
                
            html_content += """
        </main>
    </div>
</body>
</html>"""
            
            # 写入文件
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
        except Exception as e:
            raise Exception(f"HTML导出失败: {e}")
            
    def get_html_styles(self):
        """获取HTML样式"""
        return """
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Microsoft YaHei', 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.7; background: #f7fafc; color: #0f172a; }
        .container { max-width: 1000px; margin: 24px auto; padding: 0 20px; }

        .header { background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%); color: #fff; padding: 28px; border-radius: 16px; text-align: center; box-shadow: 0 10px 25px rgba(79,70,229,0.25); }
        .header h1 { font-size: 28px; margin-bottom: 8px; letter-spacing: 0.5px; }
        .meta { opacity: 0.95; }
        .meta span { margin: 0 12px; font-size: 14px; }

        .group { margin: 22px 0 12px; }
        .group h2 { background: linear-gradient(135deg, #2563eb, #3b82f6); color: #fff; padding: 12px 16px; border-radius: 10px; font-size: 18px; box-shadow: 0 6px 16px rgba(37,99,235,0.25); }

        .card { background: #fff; border-radius: 14px; box-shadow: 0 10px 30px rgba(2,8,23,0.06); border: 1px solid #e2e8f0; }
        .question.card { padding: 20px; margin: 16px 0; }
        .question-header { font-size: 16px; font-weight: 700; margin-bottom: 14px; color: #1e40af; }

        .options { margin: 12px 0; }
        .option { padding: 10px 14px; margin: 8px 0; background: #f8fafc; border-radius: 8px; border: 1px solid #e2e8f0; }

        .answer-section { display: flex; gap: 12px; margin: 12px 0; flex-wrap: wrap; }
        .answer-item { padding: 6px 12px; border-radius: 9999px; font-weight: 600; font-size: 12px; border: 1px solid transparent; }
        .my-answer { background: #eff6ff; color: #1d4ed8; border-color: #bfdbfe; }
        .correct-answer { background: #ecfdf5; color: #047857; border-color: #bbf7d0; }
        .score { background: #fff7ed; color: #9a3412; border-color: #fed7aa; }

        .explanation { background: #f0f9ff; border-left: 4px solid #0ea5e9; padding: 14px; margin: 12px 0; border-radius: 8px; }
        .explanation strong { color: #0c4a6e; }

        .statistics.card { padding: 22px; margin: 28px 0; }
        .stats-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 14px; margin-top: 12px; }
        .stat-item { text-align: center; padding: 14px; background: #f8fafc; border-radius: 10px; border: 1px solid #e2e8f0; }
        .stat-value { font-size: 24px; font-weight: 800; color: #2563eb; }
        .stat-label { color: #64748b; font-size: 12px; }
        """
        
    def format_question_html(self, question, index, config):
        """格式化单个题目为HTML"""
        question_text = question.get('title', question.get('question_text', ''))
        question_type = question.get('question_type', '')
        
        html = f'<div class="question">'
        
        # 题目标题
        if question_type:
            html += f'<div class="question-header">{index}. ({question_type}){question_text}</div>'
        else:
            html += f'<div class="question-header">{index}. {question_text}</div>'
        
        # 题干图片
        if config.get('include_images', True):
            for img in (question.get('title_images') or []):
                img_src = self._resolve_img_src(img)
                if img_src:
                    alt = img.get('alt', '图片')
                    html += f'<div class="options" style="text-align:center;"><img style="max-width: 600px; max-height: 400px;" src="{img_src}" alt="{alt}"/></div>'
            
        # 选项
        if config.get('include_options', True) and question.get('options'):
            html += '<div class="options">'
            for option in question['options']:
                label = option.get('label', '')
                content = option.get('content', '')
                opt_images = option.get('images') or []
                
                # 如果有图片，清理占位符文本
                if opt_images:
                    content = re.sub(r'\[图片选项[：:]\s*\d+张\]', '', content).strip()
                
                html += f'<div class="option">{label}. {content}</div>'
                if config.get('include_images', True) and opt_images:
                    for img in opt_images:
                        img_src = self._resolve_img_src(img)
                        if img_src:
                            alt = img.get('alt', '图片')
                            html += f'<div style="text-align:center; margin:8px 0;"><img style="max-width: 500px; max-height: 350px;" src="{img_src}" alt="{alt}"/></div>'
            html += '</div>'
            
        # 答案部分
        answer_parts = []
        if config.get('include_answers', True):
            my_answer = question.get('my_answer', '')
            correct_answer = question.get('correct_answer', '')
            if my_answer:
                answer_parts.append(f'<span class="answer-item my-answer">我的答案: {my_answer}</span>')
            if correct_answer:
                answer_parts.append(f'<span class="answer-item correct-answer">正确答案: {correct_answer}</span>')
                
        score = question.get('score', '')
        if score:
            answer_parts.append(f'<span class="answer-item score">得分: {score}</span>')
            
        if answer_parts:
            html += f'<div class="answer-section">{"".join(answer_parts)}</div>'
            
        # 解析
        if config.get('include_explanations', True):
            explanation = question.get('explanation', '')
            if explanation:
                html += f'<div class="explanation"><strong>解析:</strong> {explanation}</div>'
                if config.get('include_images', True):
                    for img in (question.get('explanation_images') or []):
                        img_src = self._resolve_img_src(img)
                        if img_src:
                            alt = img.get('alt', '图片')
                            html += f'<div style="text-align:center; margin:8px 0;"><img style="max-width: 600px; max-height: 400px;" src="{img_src}" alt="{alt}"/></div>'
                
        html += '</div>'
        return html

    def export_pdf(self, questions_data, filepath, config):
        """导出为PDF格式（简洁版，靠近用户脚本效果）"""
        if not PDF_AVAILABLE:
            raise Exception("PDF导出不可用，请安装 reportlab")
        try:
            # 准备文档
            page_size = A4
            doc = SimpleDocTemplate(
                filepath,
                pagesize=page_size,
                leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36
            )
            styles = getSampleStyleSheet()
            story = []

            title_text = config.get('custom_title') or "题目导出"
            story.append(Paragraph(title_text, styles['Title']))
            story.append(Spacer(1, 12))
            meta = f"导出时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}  题目总数: {len(questions_data)}"
            story.append(Paragraph(meta, styles['Normal']))
            story.append(Spacer(1, 18))

            grouped_questions = self.group_questions(questions_data, config)

            for group_name, questions in grouped_questions.items():
                if group_name != 'default':
                    story.append(Paragraph(f"{group_name} ({len(questions)}题)", styles['Heading2']))
                    story.append(Spacer(1, 8))

                for idx, q in enumerate(questions, 1):
                    q_text = get_question_field(q, 'content', '')
                    q_type = get_question_field(q, 'question_type', '')
                    header = f"{idx}. ({q_type}){q_text}" if q_type else f"{idx}. {q_text}"
                    story.append(Paragraph(header, styles['Heading4']))
                    story.append(Spacer(1, 6))

                    # 题干图片
                    if config.get('include_images', True):
                        for img in get_question_field(q, 'content_images', []):
                            img_src = self._resolve_img_src(img)
                            rl_img = self._image_from_base64_for_pdf(img_src, max_width=460)
                            if rl_img:
                                story.append(rl_img)
                                story.append(Spacer(1, 6))

                    # 选项
                    if config.get('include_options', True) and q.get('options'):
                        for opt in q['options']:
                            label = opt.get('label', '')
                            content = opt.get('content', '')
                            opt_images = opt.get('images') or []
                            
                            # 如果有图片，清理占位符文本
                            if opt_images:
                                content = re.sub(r'\[图片选项[：:]\s*\d+张\]', '', content).strip()
                            
                            story.append(Paragraph(f"{label}. {content}", styles['Normal']))
                            if config.get('include_images', True) and opt_images:
                                for img in opt_images:
                                    img_src = self._resolve_img_src(img)
                                    rl_img = self._image_from_base64_for_pdf(img_src, max_width=420)
                                    if rl_img:
                                        story.append(rl_img)
                            story.append(Spacer(1, 4))

                    # 答案
                    if config.get('include_answers', True):
                        my_answer = get_question_field(q, 'my_answer', '')
                        correct_answer = get_question_field(q, 'correct_answer', '')
                        if my_answer:
                            story.append(Paragraph(f"我的答案: {my_answer}", styles['Normal']))
                        if correct_answer:
                            story.append(Paragraph(f"正确答案: {correct_answer}", styles['Normal']))

                    # 得分
                    if q.get('score'):
                        story.append(Paragraph(f"得分: {q.get('score')}", styles['Normal']))

                    # 解析
                    if config.get('include_explanations', True) and q.get('explanation'):
                        story.append(Spacer(1, 4))
                        story.append(Paragraph(f"解析: {q.get('explanation')}", styles['Normal']))
                        if config.get('include_images', True):
                            for img in (q.get('explanation_images') or []):
                                img_src = self._resolve_img_src(img)
                                rl_img = self._image_from_base64_for_pdf(img_src, max_width=460)
                                if rl_img:
                                    story.append(rl_img)

                    story.append(Spacer(1, 10))

            # 统计
            if config.get('include_statistics', True):
                stats = self.calculate_statistics(questions_data)
                story.append(Spacer(1, 12))
                story.append(Paragraph("统计信息", styles['Heading2']))
                story.append(Paragraph(f"总题目数: {stats.get('total_questions', 0)}", styles['Normal']))
                story.append(Paragraph(f"正确题目: {stats.get('correct_count', 0)}", styles['Normal']))
                story.append(Paragraph(f"正确率: {stats.get('correct_rate', 0):.1f}%", styles['Normal']))

            doc.build(story)
        except Exception as e:
            raise Exception(f"PDF导出失败: {e}")

    def _image_from_base64_for_pdf(self, src: str, max_width: int = 460):
        """将base64或URL转为reportlab Image（目前仅处理base64）"""
        try:
            if not src:
                return None
            if src.startswith('data:image'):
                b64 = src.split(',')[1]
                data = base64.b64decode(b64)
                bio = BytesIO(data)
                img = RLImage(bio)
                # 调整宽度
                if img.drawWidth > max_width:
                    scale = max_width / float(img.drawWidth)
                    img.drawWidth = max_width
                    img.drawHeight = img.drawHeight * scale
                return img
            # 如需支持URL，这里可扩展请求后转BytesIO
            return None
        except Exception:
            return None
        
    def export_json(self, questions_data, filepath, config):
        """导出为JSON格式"""
        try:
            # 准备导出数据
            export_data = {
                'metadata': {
                    'title': config.get('custom_title') or '题目导出',
                    'export_time': datetime.now().isoformat(),
                    'total_questions': len(questions_data),
                    'export_config': config,
                    'version': '2.0'
                },
                'questions': questions_data
            }
            
            # 添加统计信息
            if config.get('include_statistics', True):
                export_data['statistics'] = self.calculate_statistics(questions_data)
                
            # 写入文件
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)
                
        except Exception as e:
            raise Exception(f"JSON导出失败: {e}")
            
    def export_docx(self, questions_data, filepath, config):
        """导出为Word文档格式"""
        if not DOCX_AVAILABLE:
            raise Exception("Word导出不可用，请安装 python-docx")
            
        try:
            doc = Document()
            
            # 设置文档标题
            title = config.get('custom_title') or "题目导出"
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元信息
            meta_paragraph = doc.add_paragraph()
            meta_paragraph.add_run(f"导出时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
            meta_paragraph.add_run(f"题目总数: {len(questions_data)}")
            meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # 按配置分组
            grouped_questions = self.group_questions(questions_data, config)
            
            for group_name, questions in grouped_questions.items():
                if group_name != "default":
                    doc.add_heading(f"{group_name} ({len(questions)}题)", 1)
                    
                for i, question in enumerate(questions, 1):
                    self.add_question_to_docx(doc, question, i, config)
                    # 大文档分段保存，降低丢失风险与内存峰值
                    if i % 50 == 0:
                        try:
                            doc.save(filepath)
                        except Exception:
                            pass
                    
            # 保存文档
            doc.save(filepath)
            
        except Exception as e:
            raise Exception(f"Word导出失败: {e}")
            
    def add_question_to_docx(self, doc, question, index, config):
        """添加题目到Word文档"""
        question_text = question.get('title', question.get('question_text', ''))
        question_type = question.get('question_type', '')
        
        # 题目标题
        if question_type:
            heading = doc.add_heading(f"{index}. ({question_type}){question_text}", 2)
        else:
            heading = doc.add_heading(f"{index}. {question_text}", 2)
        
        # 题干图片
        if config.get('include_images', True):
            for img in (question.get('title_images') or []):
                img_src = self._resolve_img_src(img)
                try:
                    if img_src and img_src.startswith('data:image'):
                        image_bytes = base64.b64decode(img_src.split(',')[1])
                        doc.add_picture(BytesIO(image_bytes), width=Inches(5.5))
                    # 如果是URL，这里可扩展下载后再插入；当前仅处理base64
                except Exception:
                    pass
            
        # 选项
        if config.get('include_options', True) and question.get('options'):
            options_paragraph = doc.add_paragraph("选项:")
            for option in question['options']:
                label = option.get('label', '')
                content = option.get('content', '')
                opt_images = option.get('images') or []
                
                # 如果有图片，清理占位符文本
                if opt_images:
                    content = re.sub(r'\[图片选项[：:]\s*\d+张\]', '', content).strip()
                
                doc.add_paragraph(f"{label}. {content}", style='List Bullet')
                if config.get('include_images', True) and opt_images:
                    for img in opt_images:
                        img_src = self._resolve_img_src(img)
                        try:
                            if img_src and img_src.startswith('data:image'):
                                image_bytes = base64.b64decode(img_src.split(',')[1])
                                doc.add_picture(BytesIO(image_bytes), width=Inches(4.5))
                        except Exception:
                            pass
                
        # 答案
        if config.get('include_answers', True):
            answer_paragraph = doc.add_paragraph()
            my_answer = question.get('my_answer', '')
            correct_answer = question.get('correct_answer', '')
            if my_answer:
                answer_paragraph.add_run("我的答案: ").bold = True
                answer_paragraph.add_run(f"{my_answer}\n")
            if correct_answer:
                answer_paragraph.add_run("正确答案: ").bold = True
                answer_paragraph.add_run(f"{correct_answer}\n")
                
        # 得分
        score = question.get('score', '')
        if score:
            score_paragraph = doc.add_paragraph()
            score_paragraph.add_run("得分: ").bold = True
            score_paragraph.add_run(str(score))
            
        # 解析
        if config.get('include_explanations', True):
            explanation = question.get('explanation', '')
            if explanation:
                explanation_paragraph = doc.add_paragraph()
                explanation_paragraph.add_run("解析: ").bold = True
                explanation_paragraph.add_run(explanation)
                if config.get('include_images', True):
                    for img in (question.get('explanation_images') or []):
                        img_src = self._resolve_img_src(img)
                        try:
                            if img_src and img_src.startswith('data:image'):
                                image_bytes = base64.b64decode(img_src.split(',')[1])
                                doc.add_picture(BytesIO(image_bytes), width=Inches(5.5))
                        except Exception:
                            pass
                
        doc.add_paragraph()  # 空行分隔
        
    def export_excel(self, questions_data, filepath, config):
        """导出为Excel格式"""
        if not EXCEL_AVAILABLE:
            raise Exception("Excel导出不可用，请安装 xlsxwriter")
            
        try:
            # 使用xlsxwriter（如果可用）
            if XLSXWRITER_AVAILABLE:
                self._export_excel_xlsxwriter(questions_data, filepath, config)
            else:
                raise Exception("Excel导出不可用，请安装 xlsxwriter")
                
        except Exception as e:
            raise Exception(f"Excel导出失败: {e}")
            
    def _export_excel_xlsxwriter(self, questions_data, filepath, config):
        """使用xlsxwriter导出Excel"""
        if not XLSXWRITER_AVAILABLE:
            raise Exception("xlsxwriter not available")
        
        # 确保 xlsxwriter 已导入
        import xlsxwriter
        workbook = xlsxwriter.Workbook(filepath)
        worksheet = workbook.add_worksheet('题目列表')
        
        # 设置样式
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#4472C4',
            'font_color': 'white',
            'border': 1
        })
        
        cell_format = workbook.add_format({'border': 1, 'text_wrap': True})
        
        # 写入标题行
        headers = ['序号', '作业', '题目', '类型', '选项', '我的答案', '正确答案', '得分', '是否正确']
        if config.get('include_explanations', True):
            headers.append('解析')
            
        for col, header in enumerate(headers):
            worksheet.write(0, col, header, header_format)
            
        # 写入数据
        for row, question in enumerate(questions_data, 1):
            worksheet.write(row, 0, row, cell_format)  # 序号
            worksheet.write(row, 1, question.get('homework_title', ''), cell_format)  # 作业
            worksheet.write(row, 2, question.get('title', question.get('question_text', '')), cell_format)  # 题目
            worksheet.write(row, 3, question.get('question_type', ''), cell_format)  # 类型
            
            # 选项
            options_text = ""
            if question.get('options'):
                options_text = "\n".join([f"{opt.get('label', '')}.{opt.get('content', '')}" 
                                        for opt in question['options']])
            worksheet.write(row, 4, options_text, cell_format)
            
            worksheet.write(row, 5, question.get('my_answer', ''), cell_format)  # 我的答案
            worksheet.write(row, 6, question.get('correct_answer', ''), cell_format)  # 正确答案
            worksheet.write(row, 7, question.get('score', ''), cell_format)  # 得分
            worksheet.write(row, 8, '正确' if question.get('is_correct', False) else '错误', cell_format)  # 是否正确
            
            if config.get('include_explanations', True):
                worksheet.write(row, 9, question.get('explanation', ''), cell_format)  # 解析
                
        # 调整列宽
        worksheet.set_column(0, 0, 8)   # 序号
        worksheet.set_column(1, 1, 20)  # 作业
        worksheet.set_column(2, 2, 40)  # 题目
        worksheet.set_column(3, 3, 12)  # 类型
        worksheet.set_column(4, 4, 30)  # 选项
        worksheet.set_column(5, 8, 15)  # 答案列
        if config.get('include_explanations', True):
            worksheet.set_column(9, 9, 30)  # 解析
            
        workbook.close()
        
    def group_questions(self, questions_data, config):
        """按配置分组题目"""
        group_by = config.get('group_by', '不分组')
        
        if group_by == '不分组':
            return {'default': questions_data}
        elif group_by == '按作业分组':
            groups = {}
            for question in questions_data:
                homework = question.get('homework_title', '未知作业')
                if homework not in groups:
                    groups[homework] = []
                groups[homework].append(question)
            return groups
        elif group_by == '按题目类型分组':
            groups = {}
            for question in questions_data:
                qtype = question.get('question_type', '未知类型')
                if qtype not in groups:
                    groups[qtype] = []
                groups[qtype].append(question)
            return groups
        elif group_by == '按正确性分组':
            groups = {'正确题目': [], '错误题目': []}
            for question in questions_data:
                if question.get('is_correct', False):
                    groups['正确题目'].append(question)
                else:
                    groups['错误题目'].append(question)
            return groups
        else:
            return {'default': questions_data}
            
    def calculate_statistics(self, questions_data):
        """计算统计信息"""
        if not questions_data:
            return {}
            
        stats = {
            'total_questions': len(questions_data),
            'question_types': {},
            'homework_distribution': {},
            'correct_count': 0,
            'total_score': 0,
            'average_score': 0
        }
        
        total_score = 0
        scored_questions = 0
        
        for question in questions_data:
            # 题目类型统计
            qtype = question.get('question_type', '未知类型')
            stats['question_types'][qtype] = stats['question_types'].get(qtype, 0) + 1
            
            # 作业分布统计
            homework = question.get('homework_title', '未知作业')
            stats['homework_distribution'][homework] = stats['homework_distribution'].get(homework, 0) + 1
            
            # 正确性统计
            if question.get('is_correct', False):
                stats['correct_count'] += 1
                
            # 得分统计
            score = question.get('score', 0)
            if isinstance(score, (int, float)) and score > 0:
                total_score += score
                scored_questions += 1
                
        stats['total_score'] = total_score
        stats['correct_rate'] = (stats['correct_count'] / len(questions_data) * 100) if questions_data else 0
        stats['average_score'] = (total_score / scored_questions) if scored_questions > 0 else 0
        
        return stats
        
    def generate_statistics_markdown(self, questions_data):
        """生成Markdown格式的统计信息"""
        stats = self.calculate_statistics(questions_data)
        
        content = [
            "##  统计信息",
            "",
            f"- **总题目数**: {stats['total_questions']}",
            f"- **正确题目**: {stats['correct_count']}",
            f"- **正确率**: {stats['correct_rate']:.1f}%",
            f"- **总得分**: {stats['total_score']}",
            f"- **平均得分**: {stats['average_score']:.1f}",
            "",
            "### 题目类型分布",
            ""
        ]
        
        for qtype, count in stats['question_types'].items():
            content.append(f"- {qtype}: {count} 题")
            
        content.extend(["", "### 作业分布", ""])
        
        for homework, count in stats['homework_distribution'].items():
            content.append(f"- {homework}: {count} 题")
            
        return '\n'.join(content)
        
    def generate_statistics_html(self, questions_data):
        """生成HTML格式的统计信息"""
        stats = self.calculate_statistics(questions_data)
        
        html = f"""
        <section class="statistics">
            <h2> 统计信息</h2>
            <div class="stats-grid">
                <div class="stat-item">
                    <div class="stat-value">{stats['total_questions']}</div>
                    <div class="stat-label">总题目数</div>
                </div>
                <div class="stat-item">
                    <div class="stat-value">{stats['correct_count']}</div>
                    <div class="stat-label">正确题目</div>
                </div>
                <div class="stat-item">
                    <div class="stat-value">{stats['correct_rate']:.1f}%</div>
                    <div class="stat-label">正确率</div>
                </div>
                <div class="stat-item">
                    <div class="stat-value">{stats['average_score']:.1f}</div>
                    <div class="stat-label">平均得分</div>
                </div>
            </div>
        </section>
        """
        
        return html
