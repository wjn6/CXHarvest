#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
题目导出模块 - 支持多种格式导出
整合多个优秀脚本的导出逻辑，提供最强大的导出功能

支持格式：
- HTML：精美网页格式，支持打印
- JSON：结构化数据，便于二次处理
- Word (DOCX)：正式文档格式
- PDF：通用便携格式
- Markdown：纯文本标记格式

"""

import os
import json
import base64
import re
from datetime import datetime
from typing import List, Dict, Optional, Any
from pathlib import Path

from core.enterprise_logger import app_logger
from core.version import __version__, APP_NAME
from core.common import get_question_field


class ExportOptions:
    """导出选项配置类"""
    
    def __init__(self):
        # 内容选项
        self.include_my_answer = True       # 包含我的答案
        self.include_correct_answer = True   # 包含正确答案
        self.include_score = True            # 包含得分信息
        self.include_analysis = True         # 包含答案解析
        self.include_statistics = True       # 包含统计信息
        self.include_images = True           # 包含图片
        self.embed_images = False            # Word/PDF中嵌入图片（而非链接）
        
        # 格式选项
        self.include_separator = False       # 题目间添加分割线
        self.include_question_number = True  # 包含题目编号
        self.include_question_type = True    # 包含题目类型
        self.show_correct_status = True      # 显示答题正确/错误状态
        
        # 元信息
        self.include_export_time = True      # 包含导出时间
        self.include_homework_title = True   # 包含作业标题
    
    def to_dict(self) -> dict:
        """转换为字典"""
        return {
            'include_my_answer': self.include_my_answer,
            'include_correct_answer': self.include_correct_answer,
            'include_score': self.include_score,
            'include_analysis': self.include_analysis,
            'include_statistics': self.include_statistics,
            'include_images': self.include_images,
            'embed_images': self.embed_images,
            'include_separator': self.include_separator,
            'include_question_number': self.include_question_number,
            'include_question_type': self.include_question_type,
            'show_correct_status': self.show_correct_status,
            'include_export_time': self.include_export_time,
            'include_homework_title': self.include_homework_title
        }


class QuestionExporter:
    """题目导出器 - 支持多种格式"""
    
    def __init__(self, questions: List[Dict], homework_title: str = "作业题目", session=None):
        """
        初始化导出器
        
        Args:
            questions: 题目列表
            homework_title: 作业标题
            session: 可选的 requests.Session，用于下载需要认证的图片
        """
        self.questions = questions
        self.homework_title = homework_title
        self.options = ExportOptions()
        self._statistics = None
        self._session = session
    
    def set_options(self, options: ExportOptions):
        """设置导出选项"""
        self.options = options
    
    def get_statistics(self) -> Dict:
        """获取统计信息"""
        if self._statistics is None:
            self._statistics = self._calculate_statistics()
        return self._statistics
    
    def _calculate_statistics(self) -> Dict:
        """计算统计信息"""
        stats = {
            'total_questions': len(self.questions),
            'correct_count': 0,
            'wrong_count': 0,
            'unanswered_count': 0,
            'total_score': 0,
            'max_score': 0,
            'total_images': 0,
            'question_types': {}
        }
        
        for q in self.questions:
            # 统计正确/错误
            is_correct = get_question_field(q, 'is_correct', None)
            if is_correct is True:
                stats['correct_count'] += 1
            elif is_correct is False:
                stats['wrong_count'] += 1
            else:
                stats['unanswered_count'] += 1
            
            # 统计分数
            score_str = get_question_field(q, 'score', '')
            if score_str:
                try:
                    score = float(re.sub(r'[^\d.]', '', str(score_str)))
                    stats['total_score'] += score
                except (ValueError, TypeError):
                    pass
            
            # 统计题型
            q_type = get_question_field(q, 'question_type', '未知')
            stats['question_types'][q_type] = stats['question_types'].get(q_type, 0) + 1
            
            # 统计图片
            content_images = get_question_field(q, 'content_images', [])
            stats['total_images'] += len(content_images)
            for opt in q.get('options', []):
                if isinstance(opt, dict):
                    stats['total_images'] += len(opt.get('images', []))
        
        # 计算正确率
        if stats['total_questions'] > 0:
            stats['accuracy'] = f"{(stats['correct_count'] / stats['total_questions'] * 100):.1f}%"
        else:
            stats['accuracy'] = "0%"
        
        return stats
    
    def _get_question_content(self, q: Dict) -> str:
        """获取题目内容，清理图片占位符"""
        content = get_question_field(q, 'content', '')
        # 清理图片占位符文本，如 [图片:333.jpg] 或 [图片:图片]
        content = re.sub(r'\[图片[：:][^\]]*\]', '', content).strip()
        return content
    
    def _get_question_type(self, q: Dict) -> str:
        """获取题目类型"""
        return get_question_field(q, 'question_type', '未知')
    
    def _get_question_answer(self, q: Dict) -> str:
        """获取正确答案"""
        return get_question_field(q, 'correct_answer', '')
    
    def _get_my_answer(self, q: Dict) -> str:
        """获取我的答案"""
        return get_question_field(q, 'my_answer', '')
    
    def _get_analysis(self, q: Dict) -> str:
        """获取解析"""
        return get_question_field(q, 'explanation', '')
    
    def _get_options(self, q: Dict) -> List:
        """获取选项列表"""
        options = q.get('options', [])
        result = []
        for opt in options:
            if isinstance(opt, dict):
                label = opt.get('label', '')
                content = opt.get('content', '')
                result.append(f"{label}. {content}")
            else:
                result.append(str(opt))
        return result
    
    def _is_correct(self, q: Dict) -> Optional[bool]:
        """判断是否正确"""
        return get_question_field(q, 'is_correct', None)
    
    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名中的非法字符"""
        from core.common import sanitize_filename
        return sanitize_filename(filename)
    
    # ==================== HTML 导出 ====================
    
    def export_html(self, output_path: str) -> bool:
        """
        导出为HTML格式
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            是否成功
        """
        try:
            html_content = self._generate_html()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            app_logger.info(f"HTML导出成功: {output_path}")
            return True
        except Exception as e:
            app_logger.error(f"HTML导出失败: {e}")
            return False
    
    def _generate_html(self) -> str:
        """生成HTML内容"""
        stats = self.get_statistics()
        
        # HTML头部 - 简洁青色主题
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.homework_title}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Microsoft YaHei", sans-serif;
            background: #f5f5f5;
            min-height: 100vh;
            padding: 20px;
        }}
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: #fff;
            border-radius: 16px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.08);
            overflow: hidden;
        }}
        .header {{
            background: linear-gradient(135deg, #0891b2 0%, #06b6d4 100%);
            color: white;
            padding: 32px;
            text-align: center;
        }}
        .header h1 {{
            font-size: 28px;
            margin-bottom: 8px;
        }}
        .header .meta {{
            opacity: 0.9;
            font-size: 14px;
        }}
        .stats-bar {{
            display: flex;
            justify-content: space-around;
            padding: 24px;
            background: #fff;
            border-bottom: 1px solid #e5e7eb;
        }}
        .stat-item {{
            text-align: center;
        }}
        .stat-value {{
            font-size: 28px;
            font-weight: bold;
            color: #1e293b;
        }}
        .stat-label {{
            font-size: 13px;
            color: #64748b;
            margin-top: 4px;
        }}
        .stat-item.correct .stat-value {{ color: #16a34a; }}
        .stat-item.wrong .stat-value {{ color: #dc2626; }}
        .filter-bar {{
            display: flex;
            gap: 8px;
            padding: 16px 24px;
            background: #f8fafc;
            border-bottom: 1px solid #e5e7eb;
        }}
        .filter-btn {{
            padding: 8px 16px;
            border: none;
            border-radius: 6px;
            font-size: 14px;
            cursor: pointer;
            transition: all 0.2s;
            background: #e5e7eb;
            color: #64748b;
        }}
        .filter-btn.active {{
            background: #0891b2;
            color: white;
        }}
        .filter-btn:hover {{
            opacity: 0.9;
        }}
        .content {{
            padding: 24px;
        }}
        .question-card {{
            background: #fff;
            border: 1px solid #e5e7eb;
            border-radius: 12px;
            padding: 20px;
            margin-bottom: 16px;
        }}
        .question-card.hidden {{
            display: none;
        }}
        .question-header {{
            display: flex;
            align-items: center;
            margin-bottom: 12px;
            flex-wrap: wrap;
            gap: 8px;
        }}
        .question-number {{
            background: #0891b2;
            color: white;
            padding: 4px 12px;
            border-radius: 6px;
            font-size: 13px;
            font-weight: 500;
        }}
        .question-type {{
            background: #f1f5f9;
            color: #64748b;
            padding: 4px 10px;
            border-radius: 6px;
            font-size: 12px;
        }}
        .question-status {{
            margin-left: auto;
            padding: 4px 10px;
            border-radius: 6px;
            font-size: 12px;
            font-weight: 500;
        }}
        .question-status.correct {{
            background: #dcfce7;
            color: #16a34a;
        }}
        .question-status.wrong {{
            background: #fee2e2;
            color: #dc2626;
        }}
        .question-content {{
            font-size: 15px;
            line-height: 1.7;
            color: #1e293b;
            margin-bottom: 16px;
        }}
        .options-list {{
            list-style: none;
            margin-bottom: 16px;
        }}
        .options-list li {{
            padding: 10px 14px;
            margin-bottom: 8px;
            background: #f8fafc;
            border-radius: 8px;
            font-size: 14px;
            color: #475569;
        }}
        .answer-section {{
            padding: 12px 16px;
            border-radius: 8px;
            margin-top: 12px;
            font-size: 14px;
        }}
        .answer-section.my-answer {{
            background: #ecfeff;
            border-left: 4px solid #0891b2;
        }}
        .answer-section.correct-answer {{
            background: #f0fdf4;
            border-left: 4px solid #22c55e;
        }}
        .answer-section.analysis {{
            background: #fef3c7;
            border-left: 4px solid #f59e0b;
        }}
        .answer-label {{
            font-weight: 600;
            margin-right: 8px;
        }}
        .score-info {{
            margin-top: 12px;
            font-size: 13px;
            color: #64748b;
        }}
        .question-image {{
            max-width: 100%;
            height: auto;
            border-radius: 8px;
            margin: 10px 0;
        }}
        .footer {{
            text-align: center;
            padding: 20px;
            background: #f8fafc;
            color: #94a3b8;
            font-size: 12px;
        }}
        @media print {{
            body {{ background: white; padding: 0; }}
            .container {{ box-shadow: none; }}
            .question-card {{ break-inside: avoid; }}
            .filter-bar {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{self.homework_title}</h1>
"""
        
        if self.options.include_export_time:
            html += f'            <div class="meta">导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>\n'
        
        html += """        </div>
"""
        
        # 统计信息
        if self.options.include_statistics:
            html += f"""        <div class="stats-bar">
            <div class="stat-item">
                <div class="stat-value">{stats['total_questions']}</div>
                <div class="stat-label">总题数</div>
            </div>
            <div class="stat-item correct">
                <div class="stat-value">{stats['correct_count']}</div>
                <div class="stat-label">正确</div>
            </div>
            <div class="stat-item wrong">
                <div class="stat-value">{stats['wrong_count']}</div>
                <div class="stat-label">错误</div>
            </div>
            <div class="stat-item">
                <div class="stat-value">{stats['accuracy']}</div>
                <div class="stat-label">正确率</div>
            </div>
        </div>
"""
        
        # 过滤按钮栏
        html += """        <div class="filter-bar">
            <button class="filter-btn active" onclick="filterQuestions('all')">全部题目</button>
            <button class="filter-btn" onclick="filterQuestions('correct')">答对题目</button>
            <button class="filter-btn" onclick="filterQuestions('wrong')">答错题目</button>
        </div>
        
        <div class="content">
"""
        
        # 题目列表
        for i, q in enumerate(self.questions, 1):
            q_type = self._get_question_type(q)
            content = self._get_question_content(q)
            options = self._get_options(q)
            is_correct = self._is_correct(q)
            
            data_correct = "true" if is_correct else "false"
            html += f'            <div class="question-card" data-correct="{data_correct}">\n'
            html += '                <div class="question-header">\n'
            
            if self.options.include_question_number:
                html += f'                    <span class="question-number">第 {i} 题</span>\n'
            
            if self.options.include_question_type:
                html += f'                    <span class="question-type">{q_type}</span>\n'
            
            if self.options.show_correct_status and is_correct is not None:
                status_class = "correct" if is_correct else "wrong"
                status_text = "✓ 正确" if is_correct else "✗ 错误"
                html += f'                    <span class="question-status {status_class}">{status_text}</span>\n'
            
            html += '                </div>\n'
            html += f'                <div class="question-content">{self._escape_html(content)}</div>\n'
            
            # 题目图片
            content_images = get_question_field(q, 'content_images', [])
            if content_images:
                html += f'                {self._render_images_html(content_images)}'
            
            # 选项（包含图片）
            raw_options = q.get('options', [])
            if raw_options:
                html += '                <ul class="options-list">\n'
                for opt in raw_options:
                    if isinstance(opt, dict):
                        label = opt.get('label', '')
                        content = opt.get('content', '')
                        opt_images = opt.get('images', [])
                        
                        # 如果有图片，清理占位符文本
                        if opt_images:
                            content = re.sub(r'\[图片选项[：:]\s*\d+张\]', '', content).strip()
                        
                        html += f'                    <li>{self._escape_html(f"{label}. {content}")}'
                        if opt_images:
                            html += self._render_images_html(opt_images)
                        html += '</li>\n'
                    else:
                        html += f'                    <li>{self._escape_html(str(opt))}</li>\n'
                html += '                </ul>\n'
            
            # 我的答案
            if self.options.include_my_answer:
                my_answer = self._get_my_answer(q)
                my_answer_images = get_question_field(q, 'my_answer_images', [])
                if my_answer or my_answer_images:
                    html += f'                <div class="answer-section my-answer"><span class="answer-label">我的答案:</span>'
                    if my_answer and '[图片' not in my_answer:
                        html += self._escape_html(my_answer)
                    html += '</div>\n'
                    if my_answer_images:
                        html += f'                {self._render_images_html(my_answer_images)}'
            
            # 正确答案
            if self.options.include_correct_answer:
                correct_answer = self._get_question_answer(q)
                correct_answer_images = get_question_field(q, 'correct_answer_images', [])
                if correct_answer or correct_answer_images:
                    html += f'                <div class="answer-section correct-answer"><span class="answer-label">正确答案:</span>'
                    if correct_answer and '[图片' not in correct_answer:
                        html += self._escape_html(correct_answer)
                    html += '</div>\n'
                    if correct_answer_images:
                        html += f'                {self._render_images_html(correct_answer_images)}'
            
            # 解析
            if self.options.include_analysis:
                analysis = self._get_analysis(q)
                if analysis:
                    html += f'                <div class="answer-section analysis"><span class="answer-label">解析:</span>{self._escape_html(analysis)}</div>\n'
            
            # 得分
            if self.options.include_score:
                score = get_question_field(q, 'score', '')
                if score:
                    html += f'                <div class="score-info">得分: {score}</div>\n'
            
            html += '            </div>\n'
            
            # 分割线
            if self.options.include_separator and i < len(self.questions):
                html += '            <hr style="border: none; border-top: 1px solid #e0e0e0; margin: 20px 0;">\n'
        
        html += f"""        </div>
        <div class="footer">
            Generated by {APP_NAME} v{__version__}
        </div>
    </div>
    
    <script>
        function filterQuestions(type) {{
            const cards = document.querySelectorAll('.question-card');
            const btns = document.querySelectorAll('.filter-btn');
            
            btns.forEach(btn => btn.classList.remove('active'));
            event.target.classList.add('active');
            
            cards.forEach(card => {{
                const isCorrect = card.dataset.correct === 'true';
                if (type === 'all') {{
                    card.classList.remove('hidden');
                }} else if (type === 'correct') {{
                    card.classList.toggle('hidden', !isCorrect);
                }} else {{
                    card.classList.toggle('hidden', isCorrect);
                }}
            }});
        }}
    </script>
</body>
</html>
"""
        return html
    
    def _escape_html(self, text: str) -> str:
        """转义HTML特殊字符"""
        if not text:
            return ""
        return (str(text)
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))
    
    def _render_images_html(self, images: List) -> str:
        """渲染图片为HTML，保持原始尺寸"""
        if not images or not self.options.include_images:
            return ""
        
        html = ""
        for img in images:
            if isinstance(img, dict):
                # 优先使用 data（base64），因为原始URL可能需要认证
                src = img.get('data') or img.get('src', '')
                alt = img.get('alt', '图片')
                # 使用原始尺寸，如果有的话
                width = img.get('width', 0)
                height = img.get('height', 0)
            else:
                src = str(img)
                alt = '图片'
                width = height = 0
            
            if src:
                safe_alt = self._escape_html(alt)
                if width > 0 and height > 0:
                    html += f'<img class="question-image" src="{src}" alt="{safe_alt}" width="{width}" height="{height}" style="max-width:100%;" />\n'
                else:
                    html += f'<img class="question-image" src="{src}" alt="{safe_alt}" style="max-width:600px;" />\n'
        
        return html
    
    def _get_image_bytes(self, img_data: dict) -> Optional[bytes]:
        """从图片数据获取字节流"""
        try:
            # 优先使用data字段（base64编码）
            data = img_data.get('data', '')
            if data and data.startswith('data:image'):
                if ',' in data:
                    base64_str = data.split(',', 1)[1]
                    return base64.b64decode(base64_str)
            
            # 尝试src字段
            src = img_data.get('src', '')
            if src and src.startswith('data:image'):
                if ',' in src:
                    base64_str = src.split(',', 1)[1]
                    return base64.b64decode(base64_str)
            
            # URL图片 - 尝试下载
            # 优先使用认证 session（超星图片需要 cookie）；
            # 若无 session 则 fallback 到无认证请求（公开图片仍可下载，需认证的会返回 403）
            if src and src.startswith('http'):
                try:
                    if self._session:
                        response = self._session.get(src, timeout=10)
                    else:
                        import requests
                        response = requests.get(src, timeout=10)
                    if response.status_code == 200:
                        return response.content
                except Exception as e:
                    app_logger.warning(f"下载图片失败: {src}, {e}")
                    return None
            
            return None
        except Exception as e:
            app_logger.warning(f"获取图片字节流失败: {e}")
            return None
    
    def _add_images_to_word(self, doc, images: List, max_width_inches: float = 6.0):
        """向Word文档添加图片，保持原始比例"""
        from docx.shared import Inches, Pt
        import io
        
        if not images or not self.options.include_images:
            return
        
        for img in images:
            if not isinstance(img, dict):
                continue
            
            img_bytes = self._get_image_bytes(img)
            if not img_bytes:
                continue
            
            try:
                # 使用PIL转换为PNG格式，确保python-docx能识别
                from PIL import Image
                img_pil = Image.open(io.BytesIO(img_bytes))
                orig_width, orig_height = img_pil.size
                
                # 转换为RGB模式（如果是RGBA或其他模式）
                if img_pil.mode in ('RGBA', 'LA', 'P'):
                    background = Image.new('RGB', img_pil.size, (255, 255, 255))
                    if img_pil.mode == 'P':
                        img_pil = img_pil.convert('RGBA')
                    background.paste(img_pil, mask=img_pil.split()[-1] if img_pil.mode == 'RGBA' else None)
                    img_pil = background
                elif img_pil.mode != 'RGB':
                    img_pil = img_pil.convert('RGB')
                
                # 保存为PNG到内存
                png_buffer = io.BytesIO()
                img_pil.save(png_buffer, format='PNG')
                png_buffer.seek(0)
                
                # 计算尺寸：像素转英寸（96 DPI），但限制最大宽度
                width_inches = orig_width / 96.0
                if width_inches > max_width_inches:
                    width_inches = max_width_inches
                
                doc.add_picture(png_buffer, width=Inches(width_inches))
            except Exception as e:
                app_logger.warning(f"Word添加图片失败: {type(e).__name__}: {e}")
    
    def _add_images_to_pdf(self, story: List, images: List, max_width: float = 400, max_height: float = 500):
        """向PDF文档添加图片
        
        Args:
            story: PDF story列表
            images: 图片数据列表
            max_width: 最大宽度（点）
            max_height: 最大高度（点），防止图片超出页面
        """
        import io
        try:
            from reportlab.platypus import Image as RLImage
            from reportlab.lib.units import inch
        except ImportError:
            return
        
        if not images or not self.options.include_images:
            return
        
        for img in images:
            img_bytes = self._get_image_bytes(img) if isinstance(img, dict) else None
            if img_bytes:
                try:
                    img_stream = io.BytesIO(img_bytes)
                    rl_img = RLImage(img_stream)
                    
                    # 按比例缩放 - 同时限制宽度和高度
                    width_ratio = max_width / rl_img.drawWidth if rl_img.drawWidth > max_width else 1
                    height_ratio = max_height / rl_img.drawHeight if rl_img.drawHeight > max_height else 1
                    ratio = min(width_ratio, height_ratio)
                    
                    if ratio < 1:
                        rl_img.drawWidth *= ratio
                        rl_img.drawHeight *= ratio
                    
                    story.append(rl_img)
                except Exception as e:
                    app_logger.warning(f"PDF添加图片失败: {e}")
    
    # ==================== JSON 导出 ====================
    
    def export_json(self, output_path: str, pretty: bool = True) -> bool:
        """
        导出为JSON格式
        
        Args:
            output_path: 输出文件路径
            pretty: 是否格式化输出
            
        Returns:
            是否成功
        """
        try:
            data = self._generate_json_data()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                if pretty:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                else:
                    json.dump(data, f, ensure_ascii=False)
            
            app_logger.info(f"JSON导出成功: {output_path}")
            return True
        except Exception as e:
            app_logger.error(f"JSON导出失败: {e}")
            return False
    
    def _generate_json_data(self) -> Dict:
        """生成JSON数据"""
        stats = self.get_statistics() if self.options.include_statistics else None
        
        # 根据选项过滤题目数据
        filtered_questions = []
        for q in self.questions:
            filtered = {
                'content': self._get_question_content(q),
                'options': self._get_options(q),
            }

            if self.options.include_question_number:
                filtered['question_number'] = q.get('question_number', '')

            if self.options.include_question_type:
                filtered['question_type'] = self._get_question_type(q)

            if self.options.include_homework_title:
                filtered['homework_title'] = q.get('homework_title', '')

            section = q.get('section', None)
            if section:
                filtered['section'] = section
            
            if self.options.include_my_answer:
                filtered['my_answer'] = self._get_my_answer(q)
            
            if self.options.include_correct_answer:
                filtered['correct_answer'] = self._get_question_answer(q)
            
            if self.options.include_score:
                filtered['score'] = get_question_field(q, 'score', '')
                filtered['total_score'] = get_question_field(q, 'total_score', '')
                filtered['is_correct'] = self._is_correct(q)
            
            if self.options.include_analysis:
                filtered['explanation'] = self._get_analysis(q)
            
            if self.options.include_images:
                filtered['content_images'] = get_question_field(q, 'content_images', [])
                filtered['option_images'] = get_question_field(q, 'option_images', [])
                filtered['my_answer_images'] = get_question_field(q, 'my_answer_images', [])
                filtered['correct_answer_images'] = get_question_field(q, 'correct_answer_images', [])
                filtered['explanation_images'] = get_question_field(q, 'explanation_images', [])
            
            filtered_questions.append(filtered)
        
        return {
            'title': self.homework_title,
            'export_time': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'export_options': self.options.to_dict(),
            'statistics': stats,
            'questions': filtered_questions
        }
    
    # ==================== Markdown 导出 ====================
    
    def export_markdown(self, output_path: str) -> bool:
        """
        导出为Markdown格式
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            是否成功
        """
        try:
            md_content = self._generate_markdown()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            app_logger.info(f"Markdown导出成功: {output_path}")
            return True
        except Exception as e:
            app_logger.error(f"Markdown导出失败: {e}")
            return False
    
    def _generate_markdown(self) -> str:
        """生成Markdown内容"""
        stats = self.get_statistics()
        
        md = f"# {self.homework_title}\n\n"
        
        if self.options.include_export_time:
            md += f"> 导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        # 统计信息
        if self.options.include_statistics:
            md += "## 📊 统计信息\n\n"
            md += f"| 项目 | 数值 |\n"
            md += f"|------|------|\n"
            md += f"| 总题数 | {stats['total_questions']} |\n"
            md += f"| 正确 | {stats['correct_count']} |\n"
            md += f"| 错误 | {stats['wrong_count']} |\n"
            md += f"| 正确率 | {stats['accuracy']} |\n"
            
            if stats['question_types']:
                md += "\n### 题型分布\n\n"
                for q_type, count in stats['question_types'].items():
                    md += f"- {q_type}: {count} 题\n"
            
            md += "\n"
        
        md += "## 📝 题目列表\n\n"
        
        # 题目
        for i, q in enumerate(self.questions, 1):
            q_type = self._get_question_type(q)
            content = self._get_question_content(q)
            options = self._get_options(q)
            is_correct = self._is_correct(q)
            
            # 题目标题
            if self.options.include_question_number:
                md += f"### {i}. "
            else:
                md += "### "
            
            if self.options.include_question_type:
                md += f"[{q_type}] "
            
            md += f"{content}\n\n"
            
            # 正确/错误状态
            if self.options.show_correct_status and is_correct is not None:
                status = "✅ 正确" if is_correct else "❌ 错误"
                md += f"**状态**: {status}\n\n"
            
            # 选项
            if options:
                for opt in options:
                    md += f"- {opt}\n"
                md += "\n"
            
            # 我的答案
            if self.options.include_my_answer:
                my_answer = self._get_my_answer(q)
                if my_answer:
                    md += f"**我的答案**: {my_answer}\n\n"
            
            # 正确答案
            if self.options.include_correct_answer:
                correct_answer = self._get_question_answer(q)
                if correct_answer:
                    md += f"**正确答案**: {correct_answer}\n\n"
            
            # 解析
            if self.options.include_analysis:
                analysis = self._get_analysis(q)
                if analysis:
                    md += f"> **解析**: {analysis}\n\n"
            
            # 得分
            if self.options.include_score:
                score = get_question_field(q, 'score', '')
                if score:
                    md += f"*得分: {score}*\n\n"
            
            if self.options.include_separator and i < len(self.questions):
                md += "---\n\n"
        
        return md

    def export_excel(self, output_path: str) -> bool:
        try:
            import xlsxwriter
        except ImportError:
            app_logger.error("Excel导出需要安装xlsxwriter库: pip install xlsxwriter")
            return False

        workbook = None
        try:
            workbook = xlsxwriter.Workbook(output_path)
            worksheet = workbook.add_worksheet('题目列表')

            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#0891B2',
                'font_color': 'white',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter'
            })
            cell_format = workbook.add_format({
                'border': 1,
                'text_wrap': True,
                'valign': 'top'
            })

            headers = []
            if self.options.include_question_number:
                headers.append('序号')
            if self.options.include_homework_title:
                headers.append('作业')
            headers.append('题目')
            if self.options.include_question_type:
                headers.append('类型')
            headers.append('选项')
            if self.options.include_my_answer:
                headers.append('我的答案')
            if self.options.include_correct_answer:
                headers.append('正确答案')
            if self.options.include_score:
                headers.append('得分')
            if self.options.show_correct_status:
                headers.append('是否正确')
            if self.options.include_analysis:
                headers.append('解析')

            for col, header in enumerate(headers):
                worksheet.write(0, col, header, header_format)

            worksheet.freeze_panes(1, 0)

            for row, q in enumerate(self.questions, 1):
                col = 0

                if self.options.include_question_number:
                    worksheet.write(row, col, row, cell_format)
                    col += 1

                if self.options.include_homework_title:
                    worksheet.write(row, col, q.get('homework_title', ''), cell_format)
                    col += 1

                worksheet.write(row, col, self._get_question_content(q), cell_format)
                col += 1

                if self.options.include_question_type:
                    worksheet.write(row, col, self._get_question_type(q), cell_format)
                    col += 1

                options_text = ""
                raw_options = q.get('options', [])
                if raw_options:
                    opt_lines = []
                    for opt in raw_options:
                        if isinstance(opt, dict):
                            label = opt.get('label', '')
                            content = opt.get('content', '')
                            opt_images = opt.get('images', [])
                            if opt_images:
                                content = re.sub(r'\[图片选项[：:]\s*\d+张\]', '', content).strip()
                            if label:
                                opt_lines.append(f"{label}. {content}")
                            else:
                                opt_lines.append(str(content))
                        else:
                            opt_lines.append(str(opt))
                    options_text = "\n".join(opt_lines)
                worksheet.write(row, col, options_text, cell_format)
                col += 1

                if self.options.include_my_answer:
                    my_answer = self._get_my_answer(q)
                    if my_answer and '[图片' in my_answer:
                        my_answer = ''
                    worksheet.write(row, col, my_answer, cell_format)
                    col += 1

                if self.options.include_correct_answer:
                    correct_answer = self._get_question_answer(q)
                    if correct_answer and '[图片' in correct_answer:
                        correct_answer = ''
                    worksheet.write(row, col, correct_answer, cell_format)
                    col += 1

                if self.options.include_score:
                    worksheet.write(row, col, get_question_field(q, 'score', ''), cell_format)
                    col += 1

                if self.options.show_correct_status:
                    is_correct = self._is_correct(q)
                    status = '正确' if is_correct is True else ('错误' if is_correct is False else '')
                    worksheet.write(row, col, status, cell_format)
                    col += 1

                if self.options.include_analysis:
                    worksheet.write(row, col, self._get_analysis(q), cell_format)
                    col += 1

            col_widths = {
                '序号': 8,
                '作业': 20,
                '题目': 40,
                '类型': 12,
                '选项': 30,
                '我的答案': 15,
                '正确答案': 15,
                '得分': 10,
                '是否正确': 10,
                '解析': 30
            }
            for col, header in enumerate(headers):
                worksheet.set_column(col, col, col_widths.get(header, 15))

            workbook.close()
            workbook = None
            app_logger.info(f"Excel导出成功: {output_path}")
            return True

        except Exception as e:
            app_logger.error(f"Excel导出失败: {e}")
            return False
        finally:
            if workbook is not None:
                try:
                    workbook.close()
                except Exception:
                    pass
    
    # ==================== Word (DOCX) 导出 ====================
    
    def export_word(self, output_path: str) -> bool:
        """
        导出为Word格式 (DOCX)
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            是否成功
        """
        try:
            # 尝试导入python-docx
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            app_logger.error("Word导出需要安装python-docx库: pip install python-docx")
            return False
        
        try:
            doc = Document()
            stats = self.get_statistics()
            
            # 标题
            title = doc.add_heading(self.homework_title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 副标题信息（统计 + 时间）
            if self.options.include_statistics or self.options.include_export_time:
                info_para = doc.add_paragraph()
                info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                info_parts = []
                if self.options.include_statistics:
                    info_parts.append(f"共 {stats['total_questions']} 题")
                    info_parts.append(f"正确 {stats['correct_count']}")
                    info_parts.append(f"错误 {stats['wrong_count']}")
                    info_parts.append(f"正确率 {stats['accuracy']}")
                if self.options.include_export_time:
                    info_parts.append(datetime.now().strftime('%Y-%m-%d %H:%M'))
                
                info_run = info_para.add_run("  |  ".join(info_parts))
                info_run.font.size = Pt(10)
                info_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()
            
            # 题目列表
            for i, q in enumerate(self.questions, 1):
                q_type = self._get_question_type(q)
                content = self._get_question_content(q)
                options = self._get_options(q)
                is_correct = self._is_correct(q)
                
                # 题目标题（题号 + 类型 + 状态 在同一行）
                q_para = doc.add_paragraph()
                
                # 题号
                num_run = q_para.add_run(f"{i}. ")
                num_run.bold = True
                num_run.font.size = Pt(11)
                
                # 题型
                if self.options.include_question_type:
                    type_run = q_para.add_run(f"[{q_type}] ")
                    type_run.font.size = Pt(10)
                    type_run.font.color.rgb = RGBColor(100, 100, 100)
                
                # 状态标记
                if self.options.show_correct_status and is_correct is not None:
                    if is_correct:
                        status_run = q_para.add_run(" ✓")
                        status_run.font.color.rgb = RGBColor(34, 139, 34)
                    else:
                        status_run = q_para.add_run(" ✗")
                        status_run.font.color.rgb = RGBColor(220, 20, 60)
                
                # 题目内容
                content_para = doc.add_paragraph()
                content_run = content_para.add_run(content)
                content_run.font.size = Pt(11)
                
                # 题目图片
                content_images = get_question_field(q, 'content_images', [])
                if content_images:
                    self._add_images_to_word(doc, content_images, 4.0)
                
                # 选项（包含图片）
                raw_options = q.get('options', [])
                for opt in raw_options:
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Inches(0.3)
                    
                    if isinstance(opt, dict):
                        label = opt.get('label', '')
                        content = opt.get('content', '')
                        opt_images = opt.get('images', [])
                        
                        # 如果有图片，清理占位符文本
                        if opt_images:
                            content = re.sub(r'\[图片选项[：:]\s*\d+张\]', '', content).strip()
                        
                        opt_run = opt_para.add_run(f"{label}. {content}")
                        opt_run.font.size = Pt(10)
                        
                        # 添加选项图片
                        if opt_images:
                            self._add_images_to_word(doc, opt_images, 3.0)
                    else:
                        opt_run = opt_para.add_run(str(opt))
                        opt_run.font.size = Pt(10)
                
                # 答案区域
                ans_para = doc.add_paragraph()
                ans_para.paragraph_format.space_before = Pt(6)
                
                if self.options.include_my_answer:
                    my_answer = self._get_my_answer(q)
                    my_answer_images = get_question_field(q, 'my_answer_images', [])
                    # 过滤掉图片占位符文本
                    if my_answer and '[图片' in my_answer:
                        my_answer = ''
                    if my_answer:
                        my_run = ans_para.add_run(f"我的答案: {my_answer}    ")
                        my_run.font.size = Pt(10)
                        my_run.font.color.rgb = RGBColor(0, 100, 180)
                    # 嵌入我的答案图片
                    if my_answer_images:
                        label_para = doc.add_paragraph()
                        label_run = label_para.add_run("我的答案图片:")
                        label_run.font.size = Pt(9)
                        label_run.font.color.rgb = RGBColor(0, 100, 180)
                        self._add_images_to_word(doc, my_answer_images, 3.5)
                
                if self.options.include_correct_answer:
                    correct_answer = self._get_question_answer(q)
                    correct_answer_images = get_question_field(q, 'correct_answer_images', [])
                    if correct_answer and '[图片' in correct_answer:
                        correct_answer = ''
                    if correct_answer:
                        correct_run = ans_para.add_run(f"正确答案: {correct_answer}")
                        correct_run.font.size = Pt(10)
                        correct_run.font.color.rgb = RGBColor(0, 128, 0)
                    # 嵌入正确答案图片
                    if correct_answer_images:
                        label_para = doc.add_paragraph()
                        label_run = label_para.add_run("正确答案图片:")
                        label_run.font.size = Pt(9)
                        label_run.font.color.rgb = RGBColor(0, 128, 0)
                        self._add_images_to_word(doc, correct_answer_images, 3.5)
                
                # 得分
                if self.options.include_score:
                    score = get_question_field(q, 'score', '')
                    if score:
                        score_run = ans_para.add_run(f"    得分: {score}")
                        score_run.font.size = Pt(10)
                        score_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # 解析（单独一行）
                if self.options.include_analysis:
                    analysis = self._get_analysis(q)
                    if analysis:
                        analysis_para = doc.add_paragraph()
                        analysis_run = analysis_para.add_run(f"解析: {analysis}")
                        analysis_run.font.size = Pt(9)
                        analysis_run.font.color.rgb = RGBColor(150, 100, 50)
                        analysis_run.italic = True
                
                # 分割线
                if self.options.include_separator and i < len(self.questions):
                    separator_para = doc.add_paragraph()
                    separator_para.paragraph_format.space_before = Pt(10)
                    separator_para.paragraph_format.space_after = Pt(10)
                    separator_run = separator_para.add_run("─" * 60)
                    separator_run.font.size = Pt(8)
                    separator_run.font.color.rgb = RGBColor(200, 200, 200)
                else:
                    # 题目间距
                    doc.add_paragraph()
            
            doc.save(output_path)
            app_logger.info(f"Word导出成功: {output_path}")
            return True
            
        except Exception as e:
            app_logger.error(f"Word导出失败: {e}")
            return False
    
    # ==================== PDF 导出 ====================
    
    def export_pdf(self, output_path: str) -> bool:
        """
        导出为PDF格式
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            是否成功
        """
        try:
            # 尝试导入reportlab
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            app_logger.error("PDF导出需要安装reportlab库: pip install reportlab")
            return False
        
        try:
            # 注册中文字体 (尝试多种字体)
            chinese_font_registered = False
            font_paths = [
                "C:/Windows/Fonts/simhei.ttf",  # 黑体
                "C:/Windows/Fonts/simsun.ttc",  # 宋体
                "C:/Windows/Fonts/msyh.ttc",    # 微软雅黑
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        chinese_font_registered = True
                        break
                    except Exception:
                        continue
            
            if not chinese_font_registered:
                app_logger.warning("未找到中文字体，PDF可能显示乱码")
            
            doc = SimpleDocTemplate(output_path, pagesize=A4,
                                   rightMargin=2*cm, leftMargin=2*cm,
                                   topMargin=2*cm, bottomMargin=2*cm)
            
            story = []
            styles = getSampleStyleSheet()
            
            # 创建中文样式
            if chinese_font_registered:
                title_style = ParagraphStyle(
                    'ChineseTitle',
                    parent=styles['Title'],
                    fontName='ChineseFont',
                    fontSize=18,
                    alignment=1
                )
                normal_style = ParagraphStyle(
                    'ChineseNormal',
                    parent=styles['Normal'],
                    fontName='ChineseFont',
                    fontSize=11,
                    leading=16
                )
                heading_style = ParagraphStyle(
                    'ChineseHeading',
                    parent=styles['Heading2'],
                    fontName='ChineseFont',
                    fontSize=14
                )
            else:
                title_style = styles['Title']
                normal_style = styles['Normal']
                heading_style = styles['Heading2']
            
            stats = self.get_statistics()
            
            # 标题
            story.append(Paragraph(self.homework_title, title_style))
            story.append(Spacer(1, 12))
            
            # 副标题信息（统计 + 时间）
            info_parts = []
            if self.options.include_statistics:
                info_parts.append(f"共 {stats['total_questions']} 题")
                info_parts.append(f"正确 {stats['correct_count']}")
                info_parts.append(f"错误 {stats['wrong_count']}")
                info_parts.append(f"正确率 {stats['accuracy']}")
            if self.options.include_export_time:
                info_parts.append(datetime.now().strftime('%Y-%m-%d %H:%M'))
            
            if info_parts:
                info_style = ParagraphStyle(
                    'InfoStyle',
                    parent=normal_style,
                    fontSize=10,
                    textColor=colors.grey,
                    alignment=1
                )
                story.append(Paragraph("  |  ".join(info_parts), info_style))
                story.append(Spacer(1, 25))
            
            # 题目列表
            for i, q in enumerate(self.questions, 1):
                q_type = self._get_question_type(q)
                content = self._get_question_content(q)
                options = self._get_options(q)
                is_correct = self._is_correct(q)
                
                # 题目标题（题号 + 类型 + 状态）
                status_mark = ""
                if self.options.show_correct_status and is_correct is not None:
                    status_mark = " <font color='green'>✓</font>" if is_correct else " <font color='red'>✗</font>"
                
                q_header = f"<b>{i}.</b> "
                if self.options.include_question_type:
                    q_header += f"<font color='grey'>[{q_type}]</font> "
                q_header += status_mark
                
                story.append(Paragraph(q_header, normal_style))
                
                # 题目内容
                story.append(Paragraph(content, normal_style))
                
                # 题目图片
                content_images = get_question_field(q, 'content_images', [])
                if content_images:
                    self._add_images_to_pdf(story, content_images)
                
                story.append(Spacer(1, 4))
                
                # 选项（包含图片）
                raw_options = q.get('options', [])
                if raw_options:
                    for opt in raw_options:
                        if isinstance(opt, dict):
                            label = opt.get('label', '')
                            content = opt.get('content', '')
                            opt_images = opt.get('images', [])
                            
                            # 如果有图片，清理占位符文本
                            if opt_images:
                                content = re.sub(r'\[图片选项[：:]\s*\d+张\]', '', content).strip()
                            
                            story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{label}. {content}", normal_style))
                            
                            # 添加选项图片
                            if opt_images:
                                self._add_images_to_pdf(story, opt_images, max_width=350)
                        else:
                            story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{opt}", normal_style))
                    story.append(Spacer(1, 4))
                
                # 答案（合并到一行）
                ans_parts = []
                if self.options.include_my_answer:
                    my_answer = self._get_my_answer(q)
                    my_answer_images = get_question_field(q, 'my_answer_images', [])
                    if my_answer and '[图片' in my_answer:
                        my_answer = ''
                    if my_answer:
                        ans_parts.append(f"<font color='#0066CC'>我的答案: {my_answer}</font>")
                
                if self.options.include_correct_answer:
                    correct_answer = self._get_question_answer(q)
                    correct_answer_images = get_question_field(q, 'correct_answer_images', [])
                    if correct_answer and '[图片' in correct_answer:
                        correct_answer = ''
                    if correct_answer:
                        ans_parts.append(f"<font color='#008800'>正确答案: {correct_answer}</font>")
                
                if self.options.include_score:
                    score_val = get_question_field(q, 'score', '')
                    if score_val:
                        ans_parts.append(f"<font color='grey'>得分: {score_val}</font>")
                
                if ans_parts:
                    story.append(Paragraph("&nbsp;&nbsp;".join(ans_parts), normal_style))
                
                # 嵌入答案图片到PDF
                if self.options.include_my_answer:
                    my_answer_images = get_question_field(q, 'my_answer_images', [])
                    if my_answer_images:
                        story.append(Paragraph("<font color='#0066CC'>我的答案图片:</font>", normal_style))
                        self._add_images_to_pdf(story, my_answer_images)
                
                if self.options.include_correct_answer:
                    correct_answer_images = get_question_field(q, 'correct_answer_images', [])
                    if correct_answer_images:
                        story.append(Paragraph("<font color='#008800'>正确答案图片:</font>", normal_style))
                        self._add_images_to_pdf(story, correct_answer_images)
                
                # 解析
                if self.options.include_analysis:
                    analysis = self._get_analysis(q)
                    if analysis:
                        story.append(Paragraph(f"<i><font color='#996633'>解析: {analysis}</font></i>", normal_style))
                
                # 分割线或间距
                if self.options.include_separator and i < len(self.questions):
                    from reportlab.platypus import HRFlowable
                    story.append(Spacer(1, 10))
                    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
                    story.append(Spacer(1, 10))
                else:
                    story.append(Spacer(1, 18))
            
            doc.build(story)
            app_logger.info(f"PDF导出成功: {output_path}")
            return True
            
        except Exception as e:
            app_logger.error(f"PDF导出失败: {e}")
            return False
    
    # ==================== 批量导出 ====================
    
    def export_all(self, output_dir: str, base_name: str = None) -> Dict[str, bool]:
        """
        导出所有格式
        
        Args:
            output_dir: 输出目录
            base_name: 基础文件名（不含扩展名）
            
        Returns:
            各格式导出结果
        """
        if base_name is None:
            base_name = self._sanitize_filename(self.homework_title)
        
        os.makedirs(output_dir, exist_ok=True)
        
        results = {}
        
        # HTML
        html_path = os.path.join(output_dir, f"{base_name}.html")
        results['html'] = self.export_html(html_path)
        
        # JSON
        json_path = os.path.join(output_dir, f"{base_name}.json")
        results['json'] = self.export_json(json_path)
        
        # Markdown
        md_path = os.path.join(output_dir, f"{base_name}.md")
        results['markdown'] = self.export_markdown(md_path)
        
        # Word
        docx_path = os.path.join(output_dir, f"{base_name}.docx")
        results['word'] = self.export_word(docx_path)
        
        # PDF
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        results['pdf'] = self.export_pdf(pdf_path)
        
        return results


# 便捷导出函数
def quick_export(questions: List[Dict], 
                 homework_title: str,
                 output_path: str,
                 format: str = 'html',
                 include_my_answer: bool = True,
                 include_correct_answer: bool = True) -> bool:
    """
    快速导出题目
    
    Args:
        questions: 题目列表
        homework_title: 作业标题
        output_path: 输出路径
        format: 格式 ('html', 'json', 'md', 'word', 'pdf')
        include_my_answer: 包含我的答案
        include_correct_answer: 包含正确答案
        
    Returns:
        是否成功
    """
    exporter = QuestionExporter(questions, homework_title)
    exporter.options.include_my_answer = include_my_answer
    exporter.options.include_correct_answer = include_correct_answer
    
    format = format.lower()
    
    if format == 'html':
        return exporter.export_html(output_path)
    elif format == 'json':
        return exporter.export_json(output_path)
    elif format in ('md', 'markdown'):
        return exporter.export_markdown(output_path)
    elif format in ('word', 'docx'):
        return exporter.export_word(output_path)
    elif format == 'pdf':
        return exporter.export_pdf(output_path)
    else:
        app_logger.error(f"不支持的格式: {format}")
        return False
